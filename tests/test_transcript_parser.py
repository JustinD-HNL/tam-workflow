"""Tests for the transcript file parser.

Tests PDF and DOCX parsing with real in-memory files, plus edge cases
like empty files and files with no extractable text.
"""

import io

import pytest

from src.transcript.parser import parse_docx, parse_pdf


# ---------------------------------------------------------------------------
# PDF parsing tests
# ---------------------------------------------------------------------------

class TestParsePdf:
    """Test PDF text extraction."""

    def test_parse_simple_pdf(self):
        """Create a minimal PDF in memory and verify text extraction."""
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        # PdfWriter doesn't have an easy way to create pages with text,
        # so we use the reportlab-free approach: create a page and add text annotation.
        # Instead, let's use a known minimal PDF structure.
        # The simplest approach: write text via the lower-level PDF operations.

        # Alternative: use a bytes literal of a minimal valid PDF with text
        # This is a minimal valid PDF that contains the text "Hello World"
        pdf_bytes = (
            b"%PDF-1.0\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Resources<</Font<</F1 4 0 R>>>>"
            b"/Contents 5 0 R>>endobj\n"
            b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"5 0 obj<</Length 44>>\nstream\n"
            b"BT /F1 12 Tf 100 700 Td (Hello World) Tj ET\n"
            b"endstream\nendobj\n"
            b"xref\n0 6\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"0000000115 00000 n \n"
            b"0000000296 00000 n \n"
            b"0000000373 00000 n \n"
            b"trailer<</Size 6/Root 1 0 R>>\n"
            b"startxref\n469\n%%EOF"
        )

        result = parse_pdf(pdf_bytes)
        assert "Hello World" in result

    def test_parse_empty_pdf(self):
        """A PDF with no pages should return an empty string."""
        from PyPDF2 import PdfWriter

        writer = PdfWriter()
        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        result = parse_pdf(pdf_bytes)
        assert result == ""

    def test_parse_multi_page_pdf(self):
        """PDF with multiple pages should join text with double newlines."""
        # Create a 2-page PDF with text on each page
        page1_content = b"BT /F1 12 Tf 100 700 Td (Page One) Tj ET"
        page2_content = b"BT /F1 12 Tf 100 700 Td (Page Two) Tj ET"

        pdf_bytes = (
            b"%PDF-1.0\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R 6 0 R]/Count 2>>endobj\n"
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Resources<</Font<</F1 4 0 R>>>>"
            b"/Contents 5 0 R>>endobj\n"
            b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"5 0 obj<</Length 44>>\nstream\n"
            + page1_content + b"\n"
            b"endstream\nendobj\n"
            b"6 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Resources<</Font<</F1 4 0 R>>>>"
            b"/Contents 7 0 R>>endobj\n"
            b"7 0 obj<</Length 44>>\nstream\n"
            + page2_content + b"\n"
            b"endstream\nendobj\n"
            b"xref\n0 8\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"0000000115 00000 n \n"
            b"0000000296 00000 n \n"
            b"0000000373 00000 n \n"
            b"0000000469 00000 n \n"
            b"0000000650 00000 n \n"
            b"trailer<</Size 8/Root 1 0 R>>\n"
            b"startxref\n750\n%%EOF"
        )

        result = parse_pdf(pdf_bytes)
        assert "Page One" in result
        assert "Page Two" in result

    def test_parse_invalid_pdf_raises(self):
        """Non-PDF bytes should raise an exception."""
        with pytest.raises(Exception):
            parse_pdf(b"this is not a pdf file at all")


# ---------------------------------------------------------------------------
# DOCX parsing tests
# ---------------------------------------------------------------------------

class TestParseDocx:
    """Test DOCX text extraction."""

    def test_parse_simple_docx(self):
        """Create a minimal DOCX in memory and verify text extraction."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph here")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = parse_docx(docx_bytes)
        assert "Hello from DOCX" in result
        assert "Second paragraph here" in result

    def test_parse_docx_skips_empty_paragraphs(self):
        """Empty paragraphs (whitespace only) should be filtered out."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Line A")
        doc.add_paragraph("")  # empty
        doc.add_paragraph("   ")  # whitespace only
        doc.add_paragraph("Line B")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = parse_docx(docx_bytes)
        assert "Line A" in result
        assert "Line B" in result
        # The empty paragraphs should not add extra blank content
        parts = [p for p in result.split("\n\n") if p.strip()]
        assert len(parts) == 2

    def test_parse_empty_docx(self):
        """A DOCX with no paragraphs should return an empty string."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = parse_docx(docx_bytes)
        assert result == ""

    def test_parse_docx_preserves_paragraph_order(self):
        """Paragraphs should appear in document order."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First")
        doc.add_paragraph("Second")
        doc.add_paragraph("Third")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = parse_docx(docx_bytes)
        first_pos = result.index("First")
        second_pos = result.index("Second")
        third_pos = result.index("Third")
        assert first_pos < second_pos < third_pos

    def test_parse_docx_joins_with_double_newline(self):
        """Paragraphs should be joined with double newlines."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("Beta")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = parse_docx(docx_bytes)
        assert "Alpha\n\nBeta" in result

    def test_parse_invalid_docx_raises(self):
        """Non-DOCX bytes should raise an exception."""
        with pytest.raises(Exception):
            parse_docx(b"this is not a docx file at all")
